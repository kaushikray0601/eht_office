import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Technical Note: Architecting Mineral Insulated (MI) Cable Systems for High-Temperature EHT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Introduction
doc.add_heading('1. Introduction & The Physics of MI Cables', level=1)
doc.add_paragraph(
    "Mineral Insulated (MI) cables are Series Resistance heaters, built with alloy conductors and magnesium "
    "oxide insulation to withstand temperatures up to 600°C (1112°F). Unlike self-regulating cables, they output "
    "a constant wattage based on Ohm's Law and cannot be cut in the field."
)

# Math Section
doc.add_heading('Key Mathematical Formulas', level=2)

p = doc.add_paragraph()
p.add_run('1. Total Resistance (\u03A9): ').bold = True
p.add_run('R_total = R_base \u00D7 (L_pipe + L_allowances) \u00D7 C_temp')
doc.add_paragraph('Where C_temp is the temperature correction multiplier for the specific alloy (e.g., Alloy 825).')

p = doc.add_paragraph()
p.add_run('2. Power Output (Single Phase): ').bold = True
p.add_run('P_total (W) = V\u00B2 / R_total')

p = doc.add_paragraph()
p.add_run('3. Power Output (3-Phase Star/Wye): ').bold = True
p.add_run('P_total (W) = (V_line / \u221A3)\u00B2 / R_total')

p = doc.add_paragraph()
p.add_run('4. Watt Density (W/m): ').bold = True
p.add_run('W_density = P_total / L_pipe')

p = doc.add_paragraph()
p.add_run('5. Operating Current (A): ').bold = True
p.add_run('I = P_total / V_applied')

# 2. Engineering Steps
doc.add_heading('2. Engineering & Calculation Algorithm', level=1)
doc.add_paragraph("Step 1: Calculate Heat Loss (Q_loss) based on insulation and pipe size.")
doc.add_paragraph("Step 2: Select phase configuration (1-Phase loop, 3-Phase Star, Dual-Core).")
doc.add_paragraph("Step 3: Iterate through available cable resistances (Ohms/km) using the formulas above.")
doc.add_paragraph("Step 4: Filter cables where W_density < Q_loss.")
doc.add_paragraph("Step 5: Sheath Temperature Validation. Calculate Maximum Sheath Temperature (MST) using standard IEC/IEEE heat transfer equations. If MST > T-Class Limit (e.g., 200°C), discard.")

# 3. Industry Limitations
doc.add_heading('3. Subtle Aspects & Industry Practice', level=1)
doc.add_paragraph(
    "1. The As-Built Nightmare: If the pipe is built shorter than the isometric design, the factory-terminated "
    "cable must be serpentined. This increases the W/m density locally and can cause a catastrophic sheath "
    "temperature spike, violating hazardous area T-classes."
)
doc.add_paragraph(
    "2. Star-Point Imbalance: In a 3-Phase Star configuration, if one phase runs around an extra valve, its "
    "resistance changes, shifting the neutral point and causing unbalanced heating."
)

# 4. Innovations
doc.add_heading('4. Innovative Software Features', level=1)
doc.add_paragraph("1. Dynamic As-Built Tolerance Simulator: Graphically showing how the MST shifts if the pipe is \u00B15% length.")
doc.add_paragraph("2. 3-Phase Star Auto-Balancer: Automatically calculating cold lead adjustments to balance asymmetrical pipes.")

# Save the document
doc.save('/home/kr/mydev/eht_office/MI_Cable_Engineering_Note.docx')
print("DOCX generated successfully.")
